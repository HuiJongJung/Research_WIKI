from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\jinsw712\Desktop\Files\논문 발표\PhysGaussian 발표\PhysGaussian_정휘종_발표대본.docx")


slides = [
    {
        "title": "Slide 1. Title",
        "message": "오늘 발표할 논문의 큰 주제는 3D Gaussian을 렌더링 단위이자 물리 시뮬레이션 단위로 동시에 쓰는 방법임.",
        "script": [
            "안녕하세요. 오늘 발표할 논문은 CVPR 2024 논문인 PhysGaussian: Physics-Integrated 3D Gaussians for Generative Dynamics임.",
            "이 논문은 한 문장으로 말하면, 3D Gaussian Splatting으로 만든 scene에 물리 시뮬레이션을 직접 붙이는 방법임.",
            "보통 3DGS는 고품질 novel view rendering을 위한 표현으로 많이 쓰이는데, 이 논문은 거기서 한 단계 더 나아가 Gaussian 자체를 물리 particle처럼 사용함.",
            "그래서 발표의 핵심은 '보이는 Gaussian을 그대로 움직일 수 있는가'라는 질문으로 보면 됨.",
        ],
        "transition": "먼저 발표 흐름을 간단히 보고 들어가겠음.",
    },
    {
        "title": "Slide 2. Contents",
        "message": "문제의식에서 시작해 3DGS/MPM 배경을 잡고, PhysGaussian 방법과 실험, 한계를 순서대로 설명함.",
        "script": [
            "발표는 크게 네 부분으로 진행하겠음.",
            "먼저 이 논문이 왜 필요한지 motivation을 설명하고, 그 다음 배경지식으로 3D Gaussian Splatting과 MPM을 짧게 정리하겠음.",
            "그 뒤에는 deformation map과 deformation gradient를 설명한 다음, PhysGaussian의 핵심 아이디어인 What You See Is What You Simulate를 다룰 것임.",
            "마지막으로 구체적인 method detail, 실험 결과, 그리고 limitation을 짚고 마무리하겠음.",
        ],
        "transition": "그럼 먼저 이 논문이 풀려는 문제부터 보겠음.",
    },
    {
        "title": "Slide 3. Motivation",
        "message": "기존 방식은 보이는 geometry와 시뮬레이션용 geometry가 분리되어 mismatch가 생김.",
        "script": [
            "기존 novel dynamics 생성 방식은 보통 보이는 geometry를 그대로 시뮬레이션하지 않음.",
            "예를 들어 visual geometry를 mesh로 바꾸거나, object 바깥에 cage를 만들어서 그 cage를 움직이거나, 내부를 tetrahedral volume mesh로 채운 뒤 물리 계산을 수행함.",
            "이런 방식들은 모두 원래 보이는 representation과 실제 simulation에 쓰이는 representation이 다르다는 공통점이 있음.",
            "그래서 visual geometry와 simulation geometry 사이에 mismatch가 생길 수 있고, 변환 과정도 추가로 필요함.",
            "PhysGaussian은 여기서 출발해서, 보이는 representation을 그대로 simulation할 수 있는가를 질문함.",
        ],
        "transition": "이 질문을 이해하려면 먼저 3DGS가 scene을 어떻게 표현하는지 알아야 함.",
    },
    {
        "title": "Slide 4. Background: 3D Gaussian Splatting",
        "message": "3DGS는 scene을 많은 3D Gaussian primitive의 집합으로 표현함.",
        "script": [
            "3D Gaussian Splatting은 scene을 하나의 mesh나 implicit field로 표현하지 않고, 여러 개의 3D Gaussian primitive 집합으로 표현함.",
            "각 Gaussian은 center, opacity, covariance, SH coefficient 같은 값을 가짐.",
            "center는 Gaussian의 위치이고, opacity는 얼마나 불투명하게 보이는지를 나타냄.",
            "covariance는 Gaussian ellipsoid가 어느 방향으로 얼마나 퍼져 있는지를 결정함. 즉 Gaussian의 모양과 방향을 정하는 값으로 보면 됨.",
            "SH coefficient는 보는 방향에 따라 색이 달라지는 view-dependent appearance를 표현하기 위한 값임.",
        ],
        "transition": "이제 물리 시뮬레이션 쪽 배경인 MPM을 보겠음.",
    },
    {
        "title": "Slide 5. Background: Material Point Method",
        "message": "MPM은 particle 정보를 grid로 옮겨 운동량 변화와 힘 계산을 안정적으로 근사하는 방법임.",
        "script": [
            "MPM은 Material Point Method의 약자로, 물체를 particle로 표현하면서도 계산은 grid를 함께 사용하는 시뮬레이션 방법임.",
            "particle 단위로 모든 상호작용을 직접 계산하면 주변 particle 관계가 계속 바뀌어서 복잡해짐.",
            "그래서 MPM은 particle이 가진 mass, velocity, stress 같은 정보를 grid에 모은 뒤, grid 위에서 운동량 변화량, 즉 내력과 외력을 근사 계산함.",
            "그 다음 계산된 grid 결과를 다시 particle로 가져와서 position, velocity, deformation 같은 상태를 업데이트함.",
            "particle은 물질을 따라다니며 상태를 저장하기 좋고, grid는 force, collision, momentum update를 안정적으로 계산하기 좋음.",
            "PhysGaussian에서는 이 particle 자리에 3DGS primitive 자체를 사용함.",
        ],
        "transition": "다음은 PhysGaussian에서 Gaussian을 어떻게 변형할지 설명하는 핵심 개념인 deformation map과 gradient임.",
    },
    {
        "title": "Slide 6. Deformation Map / Gradient",
        "message": "deformation map은 점의 위치 이동을, deformation gradient는 그 주변의 local 변형을 설명함.",
        "script": [
            "deformation map은 material point X가 시간 t에 어디로 이동하는지를 나타내는 함수임.",
            "쉽게 말하면 어떤 점이 움직인 후 어디에 있는지를 알려주는 위치 변화 함수임.",
            "반면 deformation gradient F는 deformation map을 material position X에 대해 미분한 값임.",
            "이 값은 한 점 주변이 어떻게 회전하고, 늘어나고, shear되는지를 담고 있음.",
            "따라서 center가 어디로 이동하는지는 deformation map으로 보고, Gaussian의 local shape이 어떻게 변하는지는 deformation gradient로 본다고 이해하면 됨.",
        ],
        "transition": "이제 이 배경을 바탕으로 PhysGaussian의 핵심 아이디어를 보겠음.",
    },
    {
        "title": "Slide 7. What You See Is What You Simulate",
        "message": "PhysGaussian은 rendering primitive와 simulation particle을 통일함.",
        "script": [
            "PhysGaussian의 핵심 문구는 What You See Is What You Simulate임.",
            "즉 렌더링에서 보이는 Gaussian이 곧 시뮬레이션에서 움직이는 Gaussian이라는 의미임.",
            "이를 위해 3DGS Gaussian에 velocity, strain, stress 같은 물리 속성을 부여함.",
            "그리고 Gaussian을 MPM material particle로 사용해서 물리 법칙에 따라 움직이게 함.",
            "이렇게 하면 기존 방식처럼 mesh extraction, cage embedding, tetrahedralization 같은 중간 proxy를 만들 필요가 줄어듦.",
        ],
        "transition": "전체 파이프라인을 한 장으로 보면 다음과 같음.",
    },
    {
        "title": "Slide 8. Overview",
        "message": "static 3DGS를 만들고, 물리 시뮬레이션에 맞게 보완한 뒤, 변형된 Gaussian을 다시 렌더링함.",
        "script": [
            "전체 과정은 먼저 input image와 camera information으로 static 3DGS를 reconstruction하는 것에서 시작함.",
            "그 다음 reconstruction 과정에서 너무 길고 얇은 Gaussian이 생기지 않도록 anisotropic loss를 optional하게 사용함.",
            "또 3DGS는 visible surface 근처에 Gaussian이 몰리는 경향이 있으므로, 내부 물리 시뮬레이션을 위해 kernel filling으로 object 내부 particle을 보완함.",
            "이후 Gaussian ellipsoid를 continuum particle로 보고, CM과 MPM을 이용해 물리적으로 어떻게 움직일지 계산함.",
            "마지막으로 계산된 움직임을 Gaussian의 위치, 모양, SH 방향에 반영하고, 변형된 Gaussian을 다시 3DGS renderer로 렌더링함.",
        ],
        "transition": "이제 이 과정에서 Gaussian state가 실제로 어떻게 업데이트되는지 보겠음.",
    },
    {
        "title": "Slide 9. Physics-Integrated 3DGS",
        "message": "center는 deformation map으로, covariance는 deformation gradient로 업데이트함.",
        "script": [
            "PhysGaussian에서 가장 중요한 부분은 물리 변형을 Gaussian representation에 어떻게 반영하느냐임.",
            "먼저 Gaussian center는 deformation map phi를 이용해 업데이트함. 즉 material space에 있던 Gaussian center가 시간 t에 어디로 이동했는지를 phi가 알려줌.",
            "하지만 center만 움직이면 Gaussian ellipsoid 자체가 늘어나거나 찌그러지는 local deformation은 표현하지 못함.",
            "그래서 covariance는 deformation gradient F를 이용해 업데이트함. 수식으로는 a_p(t) = F_p(t) A_p F_p(t)^T 형태임.",
            "covariance는 Gaussian의 퍼짐과 방향을 의미하므로, 이 업데이트는 local stretch, shear, rotation을 Gaussian shape에 반영하는 과정임.",
            "다만 deformation map 전체는 비선형일 수 있어서, 각 Gaussian 주변에서는 local affine approximation으로 근사함.",
        ],
        "transition": "그 다음 문제는 Gaussian의 모양뿐 아니라 view-dependent appearance의 방향도 함께 맞춰야 한다는 점임.",
    },
    {
        "title": "Slide 10. SH Orientation",
        "message": "SH coefficient를 직접 회전시키지 않고, view direction에 inverse rotation을 적용해 같은 효과를 냄.",
        "script": [
            "3DGS에서 색은 SH coefficient와 view direction을 이용해 계산됨.",
            "문제는 Gaussian geometry가 회전했는데 SH 방향이 그대로 있으면, 물체와 appearance가 같이 회전하지 않는 것처럼 보일 수 있다는 점임.",
            "PhysGaussian은 SH coefficient나 basis 자체를 직접 바꾸지는 않음. GS framework에서 SH basis가 고정되어 있기 때문임.",
            "대신 Gaussian별 local rotation을 구한 뒤, 색을 평가할 때 view direction에 inverse rotation을 적용함.",
            "이 rotation은 deformation gradient를 polar decomposition해서 F = R S로 나눈 뒤, 그중 rotation 성분 R만 사용함.",
            "그래서 식은 f^t(d) = f^0(R^T d)처럼 표현됨.",
        ],
        "transition": "다음은 물체 내부가 비어 있는 문제를 해결하는 internal filling임.",
    },
    {
        "title": "Slide 11. Internal Filling",
        "message": "visible surface 중심의 Gaussian 분포를 보완하기 위해 내부 particle을 추가함.",
        "script": [
            "3DGS reconstruction은 기본적으로 입력 이미지에서 보이는 표면 appearance를 잘 맞추는 방향으로 최적화됨.",
            "그래서 Gaussian들이 object의 visible surface 근처에 많이 분포하고, 내부는 비어 있는 hollow shell처럼 될 수 있음.",
            "하지만 volumetric simulation을 하려면 내부 support가 필요하기 때문에, 내부가 비어 있으면 물리 시뮬레이션이 부자연스러워질 수 있음.",
            "PhysGaussian은 Gaussian opacity field를 3D grid로 이산화하고, threshold 기반 ray test로 내부 cell을 찾음.",
            "opacity가 threshold보다 낮은 영역에서 높은 영역으로 넘어가면 surface와 intersect했다고 보고, 6방향 ray와 추가 ray 교차 횟수를 이용해 internal cell을 결정함.",
            "그렇게 찾은 내부 cell에는 particle을 추가하고, nearest Gaussian의 opacity와 SH coefficient를 가져와 초기화함.",
        ],
        "transition": "내부를 채우는 것 외에도, Gaussian 모양 자체가 너무 극단적이지 않게 만드는 regularization이 있음.",
    },
    {
        "title": "Slide 12. Anisotropy Regularizer",
        "message": "너무 길고 얇은 Gaussian은 변형 시 artifact를 만들 수 있어 scale ratio를 제한함.",
        "script": [
            "3DGS에서는 anisotropic Gaussian이 static rendering quality를 높이는 데 도움이 될 수 있음.",
            "하지만 물리 시뮬레이션에서는 너무 길고 얇은 Gaussian이 큰 변형을 받을 때 surface 밖으로 튀어나오는 artifact를 만들 수 있음.",
            "그래서 PhysGaussian은 reconstruction training 중 anisotropy regularizer를 optional하게 추가함.",
            "핵심은 Gaussian의 장축과 단축 scale ratio가 threshold r을 넘지 않도록 제한하는 것임.",
            "즉 렌더링만 잘 되는 Gaussian이 아니라, 이후 deformation에서도 안정적인 Gaussian을 만들기 위한 장치로 볼 수 있음.",
        ],
        "transition": "이제 이런 방법을 적용했을 때 어떤 결과가 나오는지 실험을 보겠음.",
    },
    {
        "title": "Slide 13. Experiments",
        "message": "다양한 material behavior에 적용 가능하고, benchmark에서도 기존 방법보다 높은 PSNR을 보고함.",
        "script": [
            "실험에서는 elastic object, plastic metal, fracture, granular material, viscoplastic paste 등 여러 material behavior를 보여줌.",
            "정성 결과의 핵심은 3DGS representation을 유지하면서도 다양한 물리 기반 novel motion을 만들 수 있다는 점임.",
            "정량 결과에서는 Wolf, Stool, Plant의 bend와 twist case에서 NeRF-Editing, Deforming-NeRF, PAC-NeRF 같은 기존 방법과 비교함.",
            "Table 1을 보면 PhysGaussian이 대부분의 case에서 더 높은 PSNR을 보고함.",
            "즉 이 실험은 물리적으로 움직인 결과가 novel view rendering 품질 면에서도 경쟁력이 있다는 것을 보여주는 근거로 볼 수 있음.",
        ],
        "transition": "다만 이 방법이 모든 물리 문제를 해결하는 것은 아니고, 명확한 한계가 있음.",
    },
    {
        "title": "Slide 14. Limitations",
        "message": "lighting, material parameter, liquid, internal filling heuristic 측면에서 한계가 있음.",
        "script": [
            "첫 번째 한계는 shadow evolution 같은 복잡한 lighting 변화는 고려하지 않는다는 점임.",
            "Gaussian의 appearance 방향은 조정하지만, 실제 조명이 변하거나 그림자가 새로 생기는 현상까지 물리적으로 모델링하지는 않음.",
            "두 번째로 material parameter를 자동으로 추정하는 것이 아니라 사용자가 수동으로 조정해야 함.",
            "세 번째로 논문에서 다루는 material은 다양하지만 liquid 같은 material은 다루지 못한다고 명시되어 있음.",
            "마지막으로 internal filling은 heuristic이기 때문에 thin, open, noisy geometry에서는 내부 판별이 실패할 수 있음.",
            "따라서 PhysGaussian은 visual representation과 simulation을 잘 연결한 강한 아이디어지만, 완전한 physical reconstruction 방법이라고 보기는 어려움.",
        ],
        "transition": "마지막으로 핵심 메시지를 정리하고 발표를 마치겠음.",
    },
    {
        "title": "Slide 15. Thank You",
        "message": "PhysGaussian의 핵심은 3D Gaussian을 렌더링 primitive이자 물리 particle로 동시에 사용하는 것임.",
        "script": [
            "정리하면 PhysGaussian은 3D Gaussian을 단순한 rendering primitive가 아니라 simulation particle로도 사용하는 방법임.",
            "이를 통해 mesh, cage, tetrahedral proxy 같은 중간 변환 없이, 보이는 representation을 물리적으로 움직이는 방향을 제안함.",
            "기술적으로는 deformation map으로 center를 업데이트하고, deformation gradient로 covariance를 업데이트하며, rotation 성분으로 SH orientation을 맞춤.",
            "한계도 있지만, 3DGS를 static rendering 표현에서 physics-grounded dynamic representation으로 확장했다는 점에서 의미가 큼.",
            "이상으로 발표 마치겠음. 감사합니다.",
        ],
        "transition": "질문 받겠음.",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_bold_label_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    p.add_run(" " + text)
    for run in p.runs:
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10.5)
    return p


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[style_name]
        st.font.name = "Malgun Gothic"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("PhysGaussian 발표 대본")
    set_run_font(r, 24, True, "000000")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("대상 PPTX: PhysGaussian_정휘종.pptx")
    set_run_font(r, 10.5, False, "555555")

    add_bold_label_paragraph(
        doc,
        "사용 방법:",
        "각 슬라이드에서 핵심 메시지를 먼저 잡고, 발표 대본을 자연스럽게 말한 뒤, 넘어가는 말로 다음 슬라이드에 연결하면 됨.",
    )
    add_bold_label_paragraph(
        doc,
        "톤:",
        "슬라이드 문체는 음슴체지만, 실제 발표는 '~입니다/합니다'보다 '~임/됨'을 살짝 풀어 말하는 자연스러운 설명체로 구성함.",
    )

    for item in slides:
        doc.add_heading(item["title"], level=1)
        add_bold_label_paragraph(doc, "핵심 메시지:", item["message"])

        h = doc.add_heading("발표 대본", level=2)
        h.paragraph_format.keep_with_next = True
        for idx, sentence in enumerate(item["script"], 1):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            p.add_run(sentence)
            for run in p.runs:
                set_run_font(run, 10.5)

        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Inches(6.5)
        set_cell_margins(table)
        cell = table.cell(0, 0)
        set_cell_shading(cell, "F4F6F9")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("넘어가는 말: ")
        set_run_font(r, 10, True, "1F4D78")
        r = p.add_run(item["transition"])
        set_run_font(r, 10)
        doc.add_paragraph()

    doc.core_properties.title = "PhysGaussian 발표 대본"
    doc.core_properties.subject = "Slide-by-slide speaker script for PhysGaussian presentation"
    doc.core_properties.author = "Codex"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
