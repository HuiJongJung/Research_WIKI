from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"outputs\hybrid_surface_residual_splatting_proposal.docx"


def set_font(run, name="Malgun Gothic", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    if level == 1:
        set_font(r, size=15, bold=True, color=(46, 116, 181))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    else:
        set_font(r, size=12.5, bold=True, color=(31, 77, 120))
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(3)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_font(r, size=10)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_font(r, size=10)
    return p


def add_num(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_font(r, size=10)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.3)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    tc_pr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + " ")
    set_font(r1, size=10, bold=True, color=(31, 58, 95))
    r2 = p.add_run(text)
    set_font(r2, size=10)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Mesh/Triangle + 3DGS Hybrid Representation 연구 아이디어")
    set_font(r, size=17, bold=True, color=(11, 37, 69))

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    r = sub.add_run(
        "Stable surface는 explicit geometry가, fuzzy/view-dependent residual은 Gaussian이 담당하는 "
        "surface-residual decomposition 제안"
    )
    set_font(r, size=10, color=(85, 85, 85))

    add_heading(doc, "한 줄 요약", 1)
    add_callout(
        doc,
        "요약:",
        "3DGS가 하나의 Gaussian primitive family로 stable surface와 fuzzy appearance를 함께 설명하는 문제를 줄이기 위해, "
        "scene을 explicit mesh/triangle surface와 Gaussian residual로 분리하는 hybrid representation을 탐색한다.",
    )

    add_heading(doc, "문제의식", 1)
    add_body(
        doc,
        "3DGS는 빠르고 photorealistic한 novel view synthesis에는 강하지만, surface가 fuzzy하고 topology가 없기 때문에 "
        "normal, collision, editing, material assignment, path tracing-friendly geometry에는 한계가 있다. 반대로 mesh나 triangle은 "
        "explicit surface, topology, normal, material 정의에는 강하지만 hair, foliage, transparency, soft boundary, specular highlight, "
        "view-dependent residual처럼 하나의 명확한 표면으로 설명하기 어려운 appearance에는 약하다.",
    )
    add_body(
        doc,
        "따라서 하나의 representation으로 모든 scene component를 설명하기보다, 안정적인 surface component와 surface로 설명되지 않는 "
        "residual component를 분리하는 방향을 생각할 수 있다.",
    )

    add_heading(doc, "기존 방법과 한계", 1)
    add_heading(doc, "Effective Rank GS", 2)
    add_body(
        doc,
        "Effective Rank GS는 Gaussian covariance의 effective rank를 분석해 needle-like Gaussian을 억제하고 disk-like surface-friendly "
        "Gaussian을 유도한다. 이는 Gaussian shape regularization으로 geometry를 개선하지만, 여전히 개별 Gaussian의 shape 제어가 "
        "중심이며 explicit topology나 triangle surface를 직접 제공하지는 않는다.",
    )

    add_heading(doc, "2DGS", 2)
    add_body(
        doc,
        "2DGS는 애초에 2D Gaussian disk primitive를 사용해 surface-like representation을 강화한다. Geometry consistency에는 유리하지만, "
        "모든 scene component를 surface disk로 설명할 때 fuzzy, volumetric, view-dependent appearance를 다루는 데 한계가 있을 수 있다.",
    )

    add_heading(doc, "SuGaR", 2)
    add_body(
        doc,
        "SuGaR는 Gaussian을 surface에 align하고 mesh extraction 및 refinement를 통해 mesh-compatible representation을 얻는다. 다만 mesh와 "
        "Gaussian의 역할이 surface와 residual로 명시적으로 분해된다기보다는, Gaussian을 surface에 정렬하고 mesh를 얻는 쪽에 초점이 있다.",
    )

    add_heading(doc, "MeshGS", 2)
    add_body(
        doc,
        "MeshGS는 mesh와 Gaussian splat을 함께 사용하며 tightly-bound와 loosely-bound Gaussian을 나누어 mesh surface와 rendering detail을 "
        "결합한다. 하지만 triangle/mesh와 Gaussian이 같은 radiance를 중복 설명하는 문제를 responsibility decomposition loss로 명시적으로 "
        "다루는지는 추가 검토가 필요하다.",
    )

    add_heading(doc, "Triangle Splatting", 2)
    add_body(
        doc,
        "Triangle Splatting 계열은 Gaussian 대신 triangle primitive를 직접 최적화해 mesh-compatible novel view synthesis를 노린다. 이는 "
        "explicit surface primitive라는 장점이 있지만, Gaussian이 잘 처리하는 fuzzy detail이나 view-dependent residual을 별도 component로 "
        "유지하는 방향과는 다르다.",
    )

    add_heading(doc, "그래서 하려는 것", 1)
    add_body(doc, "제안 방향은 scene을 다음과 같이 분해하는 것이다.")
    add_callout(doc, "Representation:", "Scene = explicit mesh/triangle surface + Gaussian residual")
    add_body(
        doc,
        "Stable하고 multi-view consistent한 surface는 mesh/triangle이 담당하고, mesh로 표현하기 어려운 fuzzy detail, transparency, "
        "view-dependent appearance, uncertain region은 Gaussian residual이 담당하게 한다. 핵심은 단순히 두 representation을 같이 쓰는 것이 "
        "아니라, 어떤 region이나 primitive가 surface를 담당하고 어떤 부분이 residual로 남아야 하는지 자동으로 분리하는 것이다.",
    )

    add_heading(doc, "핵심 기여 후보", 1)
    add_bullet(doc, "Multi-view consistency, normal confidence, effective rank, residual error 등을 이용해 surface responsibility를 추정한다.")
    add_bullet(doc, "Triangle/mesh는 stable explicit surface만 담당하고, Gaussian은 triangle이 설명하지 못한 residual appearance만 담당하도록 유도한다.")
    add_bullet(doc, "Triangle과 Gaussian이 같은 영역을 중복 설명하지 않도록 overlap penalty 또는 residual-only constraint를 설계한다.")
    add_bullet(doc, "Path tracing, editing, collision, digital twin, Physical AI처럼 explicit geometry와 photorealistic detail이 동시에 필요한 응용으로 확장 가능성을 본다.")

    add_heading(doc, "실험해보고 싶은 순서", 1)
    add_num(doc, "관련 방법 분석: SuGaR, Gaussian Frosting, MeshGS, Triangle Splatting 계열이 mesh와 Gaussian의 역할을 어떻게 나누는지 정리한다.")
    add_num(doc, "평가 기준 설계: triangle과 Gaussian의 중복 표현 정도를 어떻게 측정할지 정의한다. 예를 들어 같은 pixel/depth에서 두 component가 동시에 높은 contribution을 갖는 정도를 볼 수 있다.")
    add_num(doc, "초기 baseline 구축: 3DGS 또는 Effective Rank GS를 먼저 학습하고, high-confidence surface candidate를 추출한다.")
    add_num(doc, "Responsibility assignment 실험: multi-view consistency, normal confidence, effective rank를 조합해 triangle로 승격할 영역과 Gaussian residual로 남길 영역을 나눈다.")
    add_num(doc, "Anti-overlap loss 실험: Gaussian이 triangle이 이미 설명한 stable surface를 다시 설명하지 못하도록 하는 loss를 추가하고, rendering quality와 geometry quality를 함께 확인한다.")
    add_num(doc, "어려운 scene 평가: Mip-NeRF360의 garden/tree/foliage류, Tanks and Temples, 복잡한 thin structure가 있는 scene에서 fixed surface prior 대비 이득을 검증한다.")

    add_heading(doc, "초기 판단", 1)
    add_body(
        doc,
        "이 아이디어는 GS + mesh/triangle hybrid 자체만으로는 novelty가 약할 수 있다. 이미 SuGaR, Gaussian Frosting, MeshGS처럼 가까운 "
        "방법들이 있기 때문이다. 따라서 연구의 중심은 hybrid representation이 아니라 responsibility를 자동 분리하는 기준과, 중복 표현을 줄이는 "
        "loss 및 평가 방식이 되어야 한다.",
    )
    add_body(
        doc,
        "알고리즘적 contribution이 충분히 나오지 않더라도, explicit geometry와 photorealistic residual을 함께 요구하는 digital twin, VR/AR, "
        "Physical AI 응용으로 범위를 좁히는 방향도 가능해 보인다.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
