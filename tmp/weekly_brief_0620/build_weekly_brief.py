from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUT = Path(r"C:\Users\jinsw712\Desktop\06_20_6월_3주차_진행사항.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Pt(widths[i] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=None, bold=None, color=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    style_run(run, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    style_run(run, size=10.5)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, size=11, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2, size=11)
    else:
        r = p.add_run(text)
        style_run(r, size=11)
    return p


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(3)
r = title.add_run("06 20 (6월 3주차 진행 사항)")
style_run(r, size=20, bold=True, color="0B2545")

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(14)
r = meta.add_run("작성자: jinsw712")
style_run(r, size=11, color="555555")

add_heading(doc, "1. 이번 주 진행 요약", 1)
add_bullet(doc, "기존 triangle + 3D Gaussian 동시 splatting 방향의 배경 조사를 위해 관련 surface/mesh/geometry-oriented splatting 논문을 확인함.")
add_bullet(doc, "최근 동향은 단순 novel view synthesis quality보다, reconstruction 결과를 downstream에서 실제로 사용할 수 있는 geometry로 만드는 쪽으로 이동하고 있다고 판단함.")
add_bullet(doc, "현재는 visual quality 향상 자체보다, 최종적으로 usable한 mesh를 얻기 위한 base method와 남는 한계점을 정리하는 단계임.")
add_bullet(doc, "다음 단계로는 fully connected/opaque mesh, mesh-based NVS 평가, sparse/transparent/thin 영역 한계와 관련된 논문을 추가로 읽을 예정임.")

add_heading(doc, "2. 배경 조사 논문 정리", 1)
add_para(doc, "확인 가능한 출처에서 연도, 월, venue가 보이는 경우에만 표기하였다.")

rows = [
    ["2DGS", "2024.03 / SIGGRAPH 2024", "3D Gaussian 대신 2D oriented Gaussian disk를 사용해 surface-consistent geometry를 만들려는 기법."],
    ["Effective Rank GS", "2024.09 / NeurIPS 2024", "Gaussian이 needle-like shape로 수렴하는 문제를 effective rank로 분석하고 regularization으로 geometry/normal artifact를 줄이는 기법."],
    ["MeshGS", "2024.12 / ACCV 2024", "먼저 mesh를 추출한 뒤, mesh surface에 Gaussian을 정렬/바인딩하여 high-quality rendering을 얻는 mesh-aligned GS 기법."],
    ["SuGaR", "2024 / CVPR 2024", "3DGS를 surface-aligned하게 정리한 뒤 Poisson reconstruction 등으로 editable mesh를 빠르게 추출하는 기법."],
    ["Triangle Splatting", "2025.05 / 3DV 2026", "triangle 자체를 differentiable splatting primitive로 사용해 scene을 triangle soup 형태로 직접 최적화하는 기법."],
    ["Triangle Splatting+", "2025.09 / arXiv", "shared vertex와 opaque training schedule을 도입해 standard graphics engine에 더 바로 넣을 수 있는 semi-connected triangle representation을 만드는 기법."],
]

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
headers = ["논문/방법", "확인된 연도·월 / 발표처", "짧은 설명"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, "F2F4F7")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            style_run(run, size=10, bold=True, color="0B2545")
set_repeat_table_header(table.rows[0])

for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
        for p in cells[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                style_run(run, size=9.3)

set_table_geometry(table, [1700, 2350, 5310])

add_heading(doc, "3. 조사 후 판단한 공통 한계", 1)
add_para(doc, "공통 판단: ", bold_prefix="공통 판단: ")
last = doc.paragraphs[-1]
run = last.add_run("GS+mesh 형태는 surface가 불완전하고, mesh만으로 학습하면 표현력이 부족하며, 완전한 topology를 안정적으로 표현하지 못한다고 판단함.")
style_run(run, size=11)
add_bullet(doc, "GS 기반 방법은 rendering quality는 강하지만, downstream에서 바로 쓸 수 있는 clean surface/topology를 보장하기 어렵다.")
add_bullet(doc, "Mesh 기반 방법은 graphics engine, editing, physics 등과 잘 맞지만, sparse/uncertain/fuzzy 영역에서 표현력이 떨어질 수 있다.")
add_bullet(doc, "따라서 기존의 '각 primitive의 장점을 살려 visual quality를 높이자'는 방향에서, '최종적으로 사용 가능한 mesh를 만드는 문제'로 방향을 바꾸는 것이 더 적절하다고 봄.")

add_heading(doc, "4. 현재 base로 삼고 싶은 방향: Triangle Splatting 계열", 1)
add_para(doc, "Triangle Splatting과 Triangle Splatting+는 scene reconstruction 결과를 downstream에 활용하자는 취지와 직접적으로 맞아 현재 base method로 삼기 적절하다고 판단함.")
add_para(doc, "이 계열은 triangle primitive를 미분 가능한 형태로 splatting하여 scene을 표현하고, 기존 3DGS의 pruning과 densification을 triangle representation에 맞게 수정한 방법이다.")
add_para(doc, "특히 Triangle Splatting+는 shared vertex와 opaque triangle training을 통해 standard renderer/game engine으로 넘기기 쉬운 representation을 목표로 한다.")

add_heading(doc, "5. Triangle Splatting 계열의 핵심 한계", 1)
limitations = [
    "SfM 또는 Delaunay 기반 초기화에 의존하므로 sparse하거나 부정확한 영역에서 초기 geometry가 약해질 수 있다.",
    "Sparse 영역은 학습 중 contribution이 낮게 평가되어 pruning으로 사라질 수 있고, 이로 인해 mesh coverage와 connectivity가 깨질 수 있다.",
    "Triangle Splatting+도 full connectivity가 아니라 semi-connected mesh에 가까우며, clean topology를 보장하지는 못한다.",
    "Specular, transparent, fuzzy object는 opaque triangle만으로 표현하기 어렵다.",
    "Thin structure와 view-dependent detail은 triangle-only optimization에서 artifact나 geometry loss로 이어질 가능성이 있다.",
]
for item in limitations:
    add_number(doc, item)
add_para(doc, "따라서 다음 단계에서는 위 한계가 실제로 어떤 조건에서 발생하는지 확인하고, 이를 보완할 수 있는 방법을 조사하는 데 집중할 예정임.")

add_heading(doc, "6. 차후 읽어볼 논문", 1)
future = [
    ("MeshSplatting", "connected, opaque, colored triangle mesh를 직접 만드는 최신 방법으로, fully connected mesh 가능 여부와 현재 한계를 확인하기 위한 핵심 논문."),
    ("MiLo", "complete scene geometry를 이미지 기반으로 평가하는 mesh-based NVS 기준을 만든 논문."),
    ("Radiant Triangle Soup with Soft Connectivity Forces", "triangle soup에 soft connectivity force를 줘 surface continuity를 유도하는 논문으로, connectivity regularization 관점에서 참고할 예정."),
    ("2D-SuGaR", "2DGS와 mesh refinement를 결합해 geometry accuracy를 높이는 방법으로, surface extraction/refinement 관점에서 참고할 예정."),
    ("TransparentGS / TSGS 계열", "transparent/specular 영역에서 opaque mesh 또는 standard GS가 실패하는 원인을 보기 위한 관련 논문군."),
]
for name, desc in future:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r1 = p.add_run(name + " - ")
    style_run(r1, size=10.5, bold=True)
    r2 = p.add_run(desc)
    style_run(r2, size=10.5)

add_heading(doc, "참고 출처", 1)
sources = [
    "2DGS: https://arxiv.org/abs/2403.17888; https://dl.acm.org/doi/10.1145/3641519.3657428",
    "Effective Rank GS: https://openreview.net/forum?id=EwWpAPzcay",
    "MeshGS: https://link.springer.com/chapter/10.1007/978-981-96-0969-7_16; https://arxiv.org/abs/2410.08941",
    "SuGaR: https://openaccess.thecvf.com/content/CVPR2024/html/Guedon_SuGaR_Surface-Aligned_Gaussian_Splatting_for_Efficient_3D_Mesh_Reconstruction_and_CVPR_2024_paper.html",
    "Triangle Splatting: https://arxiv.org/abs/2505.19175; https://openreview.net/forum?id=beU5y3UFrP",
    "Triangle Splatting+: https://arxiv.org/abs/2509.25122",
    "MiLo: https://arxiv.org/html/2506.24096v1; https://dl.acm.org/doi/10.1145/3763339",
    "MeshSplatting: https://arxiv.org/abs/2512.06818; https://openaccess.thecvf.com/content/CVPR2026/papers/Held_MeshSplatting_Differentiable_Rendering_with_Opaque_Meshes_CVPR_2026_paper.pdf",
]
for s in sources:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(s)
    style_run(r, size=9.2, color="555555")

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("06 20 weekly research brief")
style_run(r, size=9, color="777777")

doc.save(OUT)
print(OUT)
