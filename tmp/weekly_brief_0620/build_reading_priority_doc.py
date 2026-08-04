from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\jinsw712\Desktop\usable_geometry_논문_읽기_우선순위.docx")


def style_run(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    style_run(r)
    return p


def add_para(doc, text="", after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if text:
        r = p.add_run(text)
        style_run(r, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    style_run(r, size=10.5)
    return p


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
r = title.add_run("Usable Geometry from Splatting: 논문 읽기 우선순위")
style_run(r, size=20, bold=True, color="0B2545")

p = add_para(doc, after=10)
r = p.add_run("기준: ")
style_run(r, size=11, bold=True, color="0B2545")
r = p.add_run("Triangle/mesh-based splatting에서 downstream usable mesh를 만들 때 어디가 아직 깨지는지 확인하는 순서로 정렬하였다.")
style_run(r, size=11)

add_heading(doc, "1. 우선순위 논문 리스트", 1)

rows = [
    ["1", "MeshSplatting", "connected, opaque, colored triangle mesh를 직접 최적화해 mesh-based NVS를 수행하는 최신 핵심 논문.", "connectivity를 어디까지 보장하는지, sparse/background/thin 영역의 failure가 남는지, connected가 clean manifold나 simulation-ready mesh까지 의미하는지 확인."],
    ["2", "Triangle Splatting+", "shared vertex와 opaque triangle training으로 standard graphics engine에 넣기 쉬운 semi-connected triangle representation을 만드는 논문.", "왜 full connectivity가 아니라 semi-connected에 머무는지, pruning이 topology/coverage를 어떻게 깨는지 확인."],
    ["3", "MiLo", "GS optimization 중 mesh를 differentiably extract해 volumetric representation과 surface representation을 연결하는 논문.", "mesh-based NVS metric이 complete scene geometry와 실제 usability를 어디까지 평가하는지 확인."],
    ["4", "Triangle Splatting", "triangle 자체를 differentiable splatting primitive로 사용해 scene을 triangle soup 형태로 직접 최적화하는 논문.", "mesh-compatible과 connected mesh의 차이, triangle soup의 장점과 한계를 확인."],
    ["5", "2DGS", "3D Gaussian 대신 2D oriented Gaussian disk를 사용해 surface-consistent geometry를 만들려는 방법.", "2D primitive가 geometry consistency를 어떻게 높이는지, mesh extraction/topology 문제는 어디까지 남는지 확인."],
    ["6", "SuGaR", "3DGS를 surface-aligned하게 정리하고 Poisson reconstruction 등으로 editable mesh를 빠르게 추출하는 논문.", "GS-to-mesh extraction baseline으로 보고, post-processing 기반 mesh가 usable topology를 얼마나 보장하는지 확인."],
    ["7", "GOF", "Gaussian opacity field에서 surface/mesh를 추출하는 reconstruction 계열.", "GS representation에서 surface를 정의하는 방식과 unbounded/thin/sparse region failure를 확인."],
    ["8", "RaDe-GS", "Gaussian splatting에서 depth/normal을 rasterize해 geometry reconstruction을 개선하는 방법.", "depth/normal signal이 mesh repair나 topology 판단에 활용될 수 있는지 확인."],
    ["9", "MeshGS", "mesh representation과 Gaussian splats를 결합해 high-quality rendering을 얻는 mesh-aligned GS 방법.", "mesh가 틀렸을 때 Gaussian이 보정 layer로 남는지, 최종 geometry가 실제로 usable한지 확인."],
    ["10", "Effective Rank GS", "3DGS Gaussian의 needle-like anisotropic shape 문제를 effective rank로 분석하고 regularization하는 논문.", "Gaussian artifact가 geometry 불안정성과 어떻게 연결되는지, auxiliary Gaussian branch에 필요한 artifact 억제 근거를 확인."],
    ["11", "Radiant Triangle Soup with Soft Connectivity Forces", "triangle soup에 soft connectivity force를 넣어 surface continuity를 유도하는 방향의 논문.", "학습 중 connectivity pressure를 어떻게 줄 수 있는지, topology repair/regularization에 참고할 점을 확인."],
    ["12", "2D-SuGaR", "2DGS와 mesh refinement를 결합해 geometry accuracy를 높이는 방향.", "surface extraction 이후 mesh refinement와 depth/normal prior가 sparse region 보완에 도움이 되는지 확인."],
    ["13", "GS2Mesh", "GS의 rendering/depth information을 이용해 mesh를 추출하는 방법.", "GS가 잘 표현한 영역을 최종 mesh로 어떻게 옮기는지, auxiliary primitive 흡수 방향에 참고할 점을 확인."],
    ["14", "TransparentGS / TSGS", "transparent 또는 specular object에서 GS reconstruction을 다루는 논문군.", "opaque mesh만으로 어려운 영역의 근거로 읽고, 주 연구 범위로 둘지 장기 한계로 둘지 판단."],
]

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
headers = ["순위", "논문/계열", "간략 요약", "봐야 할 점"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, "F2F4F7")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            style_run(run, size=9.2, bold=True, color="0B2545")
repeat_header(table.rows[0])

for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
        for p in cells[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                style_run(run, size=8.2)
set_table_geometry(table, [650, 1900, 3380, 3430])

add_heading(doc, "2. 추천 읽기 루트", 1)
add_bullet(doc, "MeshSplatting -> Triangle Splatting+ -> MiLo -> Triangle Splatting")
add_bullet(doc, "이후 2DGS / SuGaR / GOF / RaDe-GS로 GS-to-mesh baseline을 정리")
add_bullet(doc, "마지막으로 MeshGS / Effective Rank GS / connectivity force / transparent 계열을 보조적으로 확인")

add_heading(doc, "3. 읽은 뒤 좁혀야 할 질문", 1)
add_para(doc, "Connected opaque mesh는 이미 가능해지고 있다. 그렇다면 sparse/SfM-undercovered region, thin structure, background 영역에서 mesh coverage와 connectivity가 깨지는 문제를 어떻게 진단하고 복구할 것인가?")

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("usable geometry reading priority")
style_run(r, size=9, color="777777")

doc.save(OUT)
print(OUT)
