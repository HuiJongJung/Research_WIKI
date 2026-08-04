from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\jinsw712\Desktop\06_20_6월_3주차_진행사항_피드백반영.docx")


def style_run(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_para(doc, text="", size=11, bold=False, color=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if text:
        r = p.add_run(text)
        style_run(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    style_run(r)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, size=10.5, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        style_run(r, size=10.5)
    else:
        r = p.add_run(text)
        style_run(r, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    style_run(r, size=10.5)
    return p


def add_table(doc, headers, rows, widths, font_size=9.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                style_run(r, size=9.4, bold=True, color="0B2545")
    repeat_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i <= 1 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    style_run(r, size=font_size)
    set_table_geometry(table, widths)
    return table


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.orientation = WD_ORIENT.PORTRAIT
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
r = title.add_run("06 20 (6월 3주차 진행 사항)")
style_run(r, size=20, bold=True, color="0B2545")

subtitle = add_para(doc, "작성자: jinsw712", size=11, color="555555", after=12)

p = add_para(doc, after=10)
r = p.add_run("선요약: ")
style_run(r, size=11, bold=True, color="0B2545")
r = p.add_run(
    "이번 주에는 GS+mesh 기반 scene reconstruction의 최근 흐름을 조사한 결과, 단순 visual quality 향상보다 downstream에서 실제로 사용할 수 있는 mesh 생성이 핵심 과제라고 판단했으며, 특히 sparse/SfM-undercovered 영역에서 mesh coverage와 connectivity가 깨지는 문제를 중심으로 Triangle Splatting 계열 및 최신 mesh-based NVS 논문을 추가로 살펴볼 예정입니다."
)
style_run(r, size=11)

add_heading(doc, "1. 이번 주 진행 요약", 1)
add_bullet(doc, "기존 triangle + 3D Gaussian 동시 splatting 방향의 배경 조사를 위해 surface/mesh/geometry-oriented splatting 논문을 확인.")
add_bullet(doc, "최근 동향은 단순 NVS quality보다, reconstruction 결과를 downstream에서 실제로 사용할 수 있는 geometry로 만드는 쪽으로 이동하고 있다고 판단.")
add_bullet(doc, "기존의 '각 primitive의 장점을 살려 visual quality를 높이자'는 방향에서, 'mesh coverage, connectivity, topology를 usable하게 만드는 문제'로 연구 초점을 조정.")
add_bullet(doc, "단, connected mesh 생성 자체는 MeshSplatting 등 최신 연구와 충돌할 수 있으므로, 다음 단계에서는 sparse/background/thin structure에서 mesh가 사라지거나 topology가 깨지는 failure mode를 좁혀서 볼 예정.")

add_heading(doc, "2. 배경 조사 논문 정리", 1)
add_para(doc, "확인 가능한 출처에서 연도, 월, venue가 보이는 경우에만 표기하였다. 기존 표의 내용을 유지하되, 연구 방향과의 연결성을 함께 정리하였다.")

paper_rows = [
    ["2DGS", "2024.03 / SIGGRAPH 2024", "3D Gaussian 대신 2D oriented Gaussian disk를 사용해 surface-consistent geometry를 만들려는 기법.", "surface consistency를 강화하지만, 최종 clean mesh/topology 보장은 별도 문제로 남음."],
    ["Effective Rank GS", "2024.09 / NeurIPS 2024", "Gaussian의 needle-like shape를 effective rank로 분석하고 regularization으로 geometry/normal artifact를 줄이는 기법.", "Gaussian shape artifact 완화에는 유효하지만, mesh connectivity나 topology를 직접 해결하지는 않음."],
    ["MeshGS", "2024.12 / ACCV 2024", "먼저 mesh를 추출한 뒤, mesh surface에 Gaussian을 정렬/바인딩하여 high-quality rendering을 얻는 mesh-aligned GS 기법.", "mesh와 GS를 결합하지만, mesh가 부정확한 영역에서 GS가 보정 layer로 남을 가능성이 있음."],
    ["SuGaR", "2024 / CVPR 2024", "3DGS를 surface-aligned하게 정리한 뒤 Poisson reconstruction 등으로 editable mesh를 빠르게 추출하는 기법.", "빠른 mesh extraction baseline이나, extraction 후 topology/coverage quality는 scene 조건에 따라 한계가 있음."],
    ["Triangle Splatting", "2025.05 / 3DV 2026", "triangle 자체를 differentiable splatting primitive로 사용해 scene을 triangle soup 형태로 직접 최적화하는 기법.", "mesh-compatible primitive를 직접 쓰지만, triangle soup라 full connectivity는 보장하지 않음."],
    ["Triangle Splatting+", "2025.09 / arXiv", "shared vertex와 opaque training schedule을 도입해 standard graphics engine에 더 바로 넣을 수 있는 semi-connected triangle representation을 만드는 기법.", "downstream 방향과 가장 가깝지만, semi-connected이며 sparse/transparent/thin 영역 한계가 남음."],
]
add_table(
    doc,
    ["논문/방법", "연도·월 / 발표처", "짧은 설명", "연구 방향과의 관련성"],
    paper_rows,
    [1450, 1900, 3480, 2530],
    font_size=8.6,
)

add_heading(doc, "3. 조사 후 판단한 공통 한계", 1)
p = add_para(doc)
r = p.add_run("공통 판단: ")
style_run(r, size=11, bold=True)
r = p.add_run("기존 방법들은 rendering-compatible surface 또는 semi-connected mesh를 제공하지만, downstream simulation/editing에 필요한 clean connectivity, coverage, manifoldness를 일관되게 보장하지는 못한다고 판단한다.")
style_run(r, size=11)
add_bullet(doc, "GS 기반 방법은 rendering quality는 강하지만, downstream에서 바로 쓸 수 있는 clean surface/topology를 보장하기 어렵다.")
add_bullet(doc, "Mesh/triangle 기반 방법은 graphics engine, editing, physics와 잘 맞지만, sparse/uncertain/fuzzy 영역에서 표현력이 떨어지고 coverage가 깨질 수 있다.")
add_bullet(doc, "따라서 앞으로는 'fully connected mesh를 새로 만든다'가 아니라, 'connected mesh 접근법이 실패하는 조건과 복구 방법'을 중심으로 문제를 좁혀야 한다.")

add_heading(doc, "4. 현재 base로 삼고 싶은 방향: Triangle Splatting 계열", 1)
add_para(doc, "Triangle Splatting과 Triangle Splatting+는 scene reconstruction 결과를 downstream에 활용하자는 취지와 직접적으로 맞아 현재 출발점으로 삼기 적절하다고 판단한다.")
add_para(doc, "이 계열은 triangle primitive를 미분 가능한 형태로 splatting하여 scene을 표현하고, 기존 3DGS의 pruning과 densification을 triangle representation에 맞게 수정한 방법이다.")
add_para(doc, "다만 MeshSplatting이 connected opaque mesh를 직접 주장하고 있으므로, 단순히 Triangle Splatting 계열을 fully connected로 확장한다는 방향은 novelty risk가 크다. 따라서 base method로는 Triangle Splatting 계열을 보되, 연구 질문은 sparse/background/thin 영역에서 mesh coverage와 connectivity가 깨지는 문제로 좁히는 것이 필요하다.")

add_heading(doc, "5. Triangle Splatting 계열의 핵심 한계", 1)
limitations = [
    "SfM 또는 Delaunay 기반 초기화에 의존하므로 sparse하거나 부정확한 영역에서 초기 geometry가 약해질 수 있다.",
    "Sparse 영역은 학습 중 contribution이 낮게 평가되어 pruning으로 사라질 수 있고, 이로 인해 mesh coverage와 connectivity가 깨질 수 있다.",
    "Triangle Splatting+도 full connectivity가 아니라 semi-connected mesh에 가까우며, clean topology를 보장하지는 못한다.",
    "Specular, transparent, fuzzy object는 opaque triangle만으로 표현하기 어렵다.",
    "Thin structure와 view-dependent detail은 triangle-only optimization에서 artifact나 geometry loss로 이어질 가능성이 있다.",
]
for item in limitations:
    add_number(doc, item)

p = add_para(doc)
r = p.add_run("우선 집중할 한계: ")
style_run(r, size=11, bold=True, color="0B2545")
r = p.add_run("위 한계 중 transparent/specular까지 한 번에 다루기보다는, 먼저 sparse/SfM-undercovered 영역과 thin/background 구조에서 mesh coverage와 connectivity가 깨지는 문제를 실험적으로 확인하는 데 집중할 예정이다.")
style_run(r, size=11)

add_heading(doc, "6. 차후 읽어볼 논문 및 읽을 질문", 1)
future_rows = [
    ["MeshSplatting", "connected, opaque, colored triangle mesh를 직접 만드는 최신 방법.", "이미 connected mesh를 얼마나 해결했는가? 남는 failure는 sparse/background/thin 영역인가?"],
    ["MiLo", "complete scene geometry를 이미지 기반으로 평가하는 mesh-based NVS 기준을 만든 논문.", "mesh-based NVS metric은 geometry usability를 어디까지 평가하고, 무엇을 놓치는가?"],
    ["Triangle Splatting+", "shared vertex와 opaque triangle schedule로 semi-connected mesh를 만드는 현재 base 후보.", "pruning/densification이 connectivity와 coverage를 어디서 깨뜨리는가?"],
    ["GOF / RaDe-GS", "GS representation에서 mesh/depth/normal을 추출하거나 정렬하는 surface reconstruction 계열.", "GS-to-mesh extraction baseline이 sparse region과 topology 문제를 어떻게 다루는가?"],
    ["2D-SuGaR / Radiant Triangle Soup", "mesh refinement 또는 soft connectivity force를 통해 surface continuity를 높이는 관련 방법.", "connectivity regularization이나 refinement를 triangle-based splatting에 어떻게 참고할 수 있는가?"],
    ["TransparentGS / TSGS 계열", "transparent/specular 영역에서 standard GS 또는 opaque mesh가 실패하는 원인을 다루는 논문군.", "이 문제는 본 연구의 주 실험으로 삼을지, 장기 한계로 남길지 판단할 근거는 무엇인가?"],
]
add_table(doc, ["논문/계열", "읽는 목적", "읽을 때 볼 질문"], future_rows, [1800, 3300, 4260], font_size=8.8)

add_heading(doc, "7. 다음 주까지 정리할 항목", 1)
next_items = [
    "MeshSplatting, MiLo, Triangle Splatting+를 우선 읽고 connectivity guarantee, failure case, evaluation metric을 비교표로 정리.",
    "각 방법이 말하는 connected, semi-connected, complete scene geometry, mesh-based NVS의 의미를 구분.",
    "Sparse/background/thin structure에서 mesh coverage가 사라지는 사례를 논문 figure 또는 직접 실험 가능한 dataset 기준으로 수집.",
    "가능하다면 connected component 수, mesh coverage, hole/boundary artifact, PSNR/LPIPS 손실을 함께 보는 평가 항목 초안 작성.",
]
for item in next_items:
    add_bullet(doc, item)

add_heading(doc, "참고 출처", 1)
sources = [
    "2DGS: https://arxiv.org/abs/2403.17888; https://dl.acm.org/doi/10.1145/3641519.3657428",
    "Effective Rank GS: https://openreview.net/forum?id=EwWpAPzcay",
    "MeshGS: https://openaccess.thecvf.com/content/ACCV2024/papers/Choi_MeshGS_Adaptive_Mesh-Aligned_Gaussian_Splatting_for_High-Quality_Rendering_ACCV_2024_paper.pdf",
    "SuGaR: https://openaccess.thecvf.com/content/CVPR2024/html/Guedon_SuGaR_Surface-Aligned_Gaussian_Splatting_for_Efficient_3D_Mesh_Reconstruction_and_CVPR_2024_paper.html",
    "Triangle Splatting: https://openreview.net/forum?id=beU5y3UFrP",
    "Triangle Splatting+: https://arxiv.org/abs/2509.25122",
    "MiLo: https://anttwo.github.io/milo/; https://dl.acm.org/doi/10.1145/3763339",
    "MeshSplatting: https://meshsplatting.github.io/; https://openaccess.thecvf.com/content/CVPR2026/html/Held_MeshSplatting_Differentiable_Rendering_with_Opaque_Meshes_CVPR_2026_paper.html",
]
for source in sources:
    add_bullet(doc, source)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("06 20 weekly research brief - revised")
style_run(r, size=9, color="777777")

doc.save(OUT)
print(OUT)
